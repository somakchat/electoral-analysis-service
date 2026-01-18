"""
Bulk Ingestion Script - Indexes ALL political data to BOTH KG and OpenSearch.

This script:
1. Processes all files from political-data folder
2. Indexes to Knowledge Graph (for entity-based retrieval)
3. Indexes to OpenSearch (for semantic/hybrid search)
4. Handles multiple file formats: CSV, XLSX, DOCX, TXT

Usage:
    cd backend
    python ../scripts/bulk_ingest_all_sources.py
"""
import os
import sys
from pathlib import Path
import re
import hashlib
from datetime import datetime
from typing import List, Dict, Any, Tuple

# Add backend to path
script_dir = Path(__file__).parent
backend_dir = script_dir.parent / "backend"
sys.path.insert(0, str(backend_dir))

from app.config import settings
from app.services.rag.political_rag import PoliticalRAGSystem
from app.services.rag.embeddings import get_embedding_service
from app.services.rag.unified_vectordb import VectorDBConfig, OpenSearchVectorDB, Document
from app.services.rag.knowledge_graph import PoliticalKnowledgeGraph
from app.services.rag.data_schema import FactWithCitation


class BulkIngester:
    """Bulk ingestion to both KG and OpenSearch."""
    
    def __init__(self, data_dir: str):
        self.data_dir = Path(data_dir)
        self.rag = None
        self.kg = None
        self.os_client = None
        self.embedder = None
        
        # Track stats
        self.stats = {
            "files_processed": 0,
            "chunks_created": 0,
            "opensearch_indexed": 0,
            "kg_facts_added": 0,
            "errors": []
        }
    
    def initialize(self):
        """Initialize all components."""
        print("=" * 60)
        print("Bulk Ingestion - Initializing Components")
        print("=" * 60)
        
        # Initialize Political RAG (includes KG)
        print("\n[1/4] Initializing Political RAG...")
        self.rag = PoliticalRAGSystem(
            data_dir=str(self.data_dir),
            index_dir=str(settings.index_dir)
        )
        self.rag.initialize()
        self.kg = self.rag.kg
        print(f"  ✓ KG loaded with {len(self.kg.constituency_profiles)} constituencies")
        
        # Initialize OpenSearch
        print("\n[2/4] Initializing OpenSearch...")
        if VectorDBConfig.is_configured():
            self.os_client = OpenSearchVectorDB()
            print(f"  ✓ OpenSearch connected: {VectorDBConfig.OPENSEARCH_ENDPOINT}")
        else:
            print("  ⚠ OpenSearch not configured, skipping")
        
        # Initialize embedder
        print("\n[3/4] Initializing Embeddings...")
        self.embedder = get_embedding_service()
        print(f"  ✓ Embedder ready")
        
        print("\n[4/4] Ready to ingest!")
        print("=" * 60)
    
    def process_all_files(self):
        """Process all files in data directory."""
        print(f"\nScanning: {self.data_dir}")
        
        # Supported extensions
        extensions = {'.csv', '.xlsx', '.xls', '.xlsm', '.docx', '.txt', '.pdf'}
        
        files = []
        for ext in extensions:
            files.extend(self.data_dir.glob(f"*{ext}"))
        
        print(f"Found {len(files)} files to process\n")
        
        for i, file_path in enumerate(files, 1):
            print(f"\n[{i}/{len(files)}] Processing: {file_path.name}")
            try:
                self._process_file(file_path)
                self.stats["files_processed"] += 1
            except Exception as e:
                print(f"  ✗ Error: {e}")
                self.stats["errors"].append(f"{file_path.name}: {str(e)}")
        
        return self.stats
    
    def _process_file(self, file_path: Path):
        """Process a single file."""
        ext = file_path.suffix.lower()
        
        # Generate doc ID
        content_hash = hashlib.md5(file_path.read_bytes()).hexdigest()[:8]
        doc_id = f"{file_path.stem}_{content_hash}"
        
        # Parse file into chunks
        if ext == '.csv':
            chunks = self._process_csv(file_path)
        elif ext in {'.xlsx', '.xls', '.xlsm'}:
            chunks = self._process_excel(file_path)
        elif ext == '.docx':
            chunks = self._process_docx(file_path)
        elif ext == '.txt':
            chunks = self._process_txt(file_path)
        else:
            print(f"  Skipping unsupported format: {ext}")
            return
        
        print(f"  Created {len(chunks)} chunks")
        self.stats["chunks_created"] += len(chunks)
        
        # Index to OpenSearch
        if self.os_client and chunks:
            os_count = self._index_to_opensearch(chunks, doc_id, file_path.name)
            print(f"  OpenSearch: {os_count} documents indexed")
            self.stats["opensearch_indexed"] += os_count
        
        # Index to Knowledge Graph
        kg_count = self._index_to_kg(chunks, doc_id, file_path.name)
        print(f"  KG: {kg_count} facts added")
        self.stats["kg_facts_added"] += kg_count
    
    def _process_csv(self, file_path: Path) -> List[Dict]:
        """Process CSV file into chunks."""
        import pandas as pd
        
        chunks = []
        try:
            df = pd.read_csv(file_path, encoding='utf-8')
        except:
            df = pd.read_csv(file_path, encoding='latin1')
        
        # Create summary chunk
        summary = f"# {file_path.name}\n\n"
        summary += f"Columns: {', '.join(df.columns.tolist())}\n"
        summary += f"Rows: {len(df)}\n\n"
        
        # Sample data
        if len(df) > 0:
            summary += "Sample data:\n"
            summary += df.head(5).to_string()
        
        chunks.append({
            "text": summary,
            "metadata": {
                "source_file": file_path.name,
                "data_type": "csv_summary",
                "row_count": len(df),
                "columns": df.columns.tolist()
            }
        })
        
        # Create row-level chunks for key data
        # Check if this is electoral data
        if any(col.lower() in ['ac_name', 'constituency', 'ac_no'] for col in df.columns):
            for _, row in df.iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                if len(row_text) > 50:  # Only meaningful rows
                    chunks.append({
                        "text": row_text,
                        "metadata": {
                            "source_file": file_path.name,
                            "data_type": "electoral_data",
                            "constituency": row.get('AC_NAME', row.get('ac_name', row.get('Constituency', ''))),
                            "district": row.get('District', row.get('district', '')),
                        }
                    })
        
        return chunks
    
    def _process_excel(self, file_path: Path) -> List[Dict]:
        """Process Excel file into chunks."""
        import pandas as pd
        
        chunks = []
        try:
            # Try reading all sheets
            xlsx = pd.ExcelFile(file_path)
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                
                if df.empty:
                    continue
                
                # Create summary chunk
                summary = f"# {file_path.name} - Sheet: {sheet_name}\n\n"
                summary += f"Columns: {', '.join(str(c) for c in df.columns.tolist())}\n"
                summary += f"Rows: {len(df)}\n\n"
                
                # Include all data as text
                for _, row in df.iterrows():
                    row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                    if row_text:
                        summary += row_text + "\n"
                
                chunks.append({
                    "text": summary[:4000],  # Limit size
                    "metadata": {
                        "source_file": file_path.name,
                        "sheet_name": sheet_name,
                        "data_type": "survey_response" if "response" in file_path.name.lower() else "excel_data",
                        "row_count": len(df)
                    }
                })
                
        except Exception as e:
            print(f"    Excel read error: {e}")
        
        return chunks
    
    def _process_docx(self, file_path: Path) -> List[Dict]:
        """Process DOCX file into chunks."""
        from docx import Document
        
        chunks = []
        doc = Document(file_path)
        
        # Collect all text
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Create chunks of ~1000 chars
        current_chunk = []
        current_len = 0
        
        for para in full_text:
            if current_len + len(para) > 1000 and current_chunk:
                chunks.append({
                    "text": "\n".join(current_chunk),
                    "metadata": {
                        "source_file": file_path.name,
                        "data_type": "document"
                    }
                })
                current_chunk = []
                current_len = 0
            
            current_chunk.append(para)
            current_len += len(para)
        
        # Add remaining
        if current_chunk:
            chunks.append({
                "text": "\n".join(current_chunk),
                "metadata": {
                    "source_file": file_path.name,
                    "data_type": "document"
                }
            })
        
        return chunks
    
    def _process_txt(self, file_path: Path) -> List[Dict]:
        """Process TXT file into chunks."""
        chunks = []
        
        try:
            text = file_path.read_text(encoding='utf-8')
        except:
            text = file_path.read_text(encoding='latin1')
        
        # Split into ~1000 char chunks
        for i in range(0, len(text), 1000):
            chunk_text = text[i:i+1000]
            if chunk_text.strip():
                chunks.append({
                    "text": chunk_text,
                    "metadata": {
                        "source_file": file_path.name,
                        "data_type": "text"
                    }
                })
        
        return chunks
    
    def _index_to_opensearch(self, chunks: List[Dict], doc_id: str, file_name: str) -> int:
        """Index chunks to OpenSearch."""
        if not self.os_client or not chunks:
            return 0
        
        # Prepare texts for embedding
        texts = [c["text"] for c in chunks]
        
        # Generate embeddings in batches
        batch_size = 50
        all_embeddings = []
        
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i+batch_size]
            embeddings = self.embedder.embed(batch)
            all_embeddings.extend(embeddings)
        
        # Create Document objects
        documents = []
        for i, (chunk, embedding) in enumerate(zip(chunks, all_embeddings)):
            doc = Document(
                doc_id=f"{doc_id}_chunk_{i}",
                text=chunk["text"],
                embedding=embedding.tolist() if hasattr(embedding, 'tolist') else list(embedding),
                metadata={
                    **chunk.get("metadata", {}),
                    "source_doc_id": doc_id,
                    "source_type": "bulk_ingestion",
                    "source_file": file_name,
                    "chunk_id": f"chunk_{i}",
                    "indexed_at": datetime.now().isoformat()
                }
            )
            documents.append(doc)
        
        # Bulk index
        indexed = self.os_client.index_documents(documents)
        return indexed
    
    def _index_to_kg(self, chunks: List[Dict], doc_id: str, file_name: str) -> int:
        """Index chunks to Knowledge Graph as facts."""
        if not self.kg or not chunks:
            return 0
        
        facts_added = 0
        known_constituencies = set(self.kg.constituency_profiles.keys())
        
        party_patterns = {
            'BJP': r'\b(bjp|bharatiya janata)\b',
            'TMC': r'\b(tmc|trinamool|aitc|mamata)\b',
            'INC': r'\b(congress|inc)\b',
            'CPM': r'\b(cpm|cpim|left front|communist)\b'
        }
        
        for i, chunk in enumerate(chunks):
            text = chunk.get("text", "")
            if not text:
                continue
            
            # Extract constituency mentions
            mentioned_constituencies = []
            for ac_name in known_constituencies:
                if ac_name.lower() in text.lower():
                    mentioned_constituencies.append(ac_name)
            
            # Extract party mentions
            parties_mentioned = []
            for party, pattern in party_patterns.items():
                if re.search(pattern, text.lower()):
                    parties_mentioned.append(party)
            
            # Determine fact type
            fact_type = "document_content"
            if any(w in text.lower() for w in ['strategy', 'recommend', 'action', 'plan', 'should']):
                fact_type = "strategy"
            elif any(w in text.lower() for w in ['predict', 'forecast', 'expect', 'projected']):
                fact_type = "prediction"
            elif any(w in text.lower() for w in ['result', 'won', 'lost', 'vote share', 'margin']):
                fact_type = "electoral_result"
            elif any(w in text.lower() for w in ['survey', 'opinion', 'response', 'poll']):
                fact_type = "survey"
            
            # Create facts for constituencies
            for ac_name in mentioned_constituencies:
                fact = FactWithCitation(
                    fact_type=fact_type,
                    entity_name=ac_name,
                    fact_text=text[:500],
                    confidence=0.85,
                    source_file=file_name,
                    source_row=i,
                    source_doc_ids=[doc_id],
                    related_entities=list(set(mentioned_constituencies + parties_mentioned))
                )
                self.kg.add_fact(fact)
                facts_added += 1
            
            # Create facts for parties
            for party in parties_mentioned:
                fact = FactWithCitation(
                    fact_type=fact_type,
                    entity_name=party,
                    fact_text=text[:500],
                    confidence=0.85,
                    source_file=file_name,
                    source_row=i,
                    source_doc_ids=[doc_id],
                    related_entities=list(set(mentioned_constituencies + parties_mentioned))
                )
                self.kg.add_fact(fact)
                facts_added += 1
            
            # If no specific entities found, create a general fact
            if not mentioned_constituencies and not parties_mentioned:
                fact = FactWithCitation(
                    fact_type=fact_type,
                    entity_name="west_bengal_politics",
                    fact_text=text[:500],
                    confidence=0.75,
                    source_file=file_name,
                    source_row=i,
                    source_doc_ids=[doc_id],
                    related_entities=[]
                )
                self.kg.add_fact(fact)
                facts_added += 1
        
        return facts_added
    
    def print_summary(self):
        """Print ingestion summary."""
        print("\n" + "=" * 60)
        print("INGESTION SUMMARY")
        print("=" * 60)
        print(f"Files Processed: {self.stats['files_processed']}")
        print(f"Chunks Created: {self.stats['chunks_created']}")
        print(f"OpenSearch Documents: {self.stats['opensearch_indexed']}")
        print(f"KG Facts Added: {self.stats['kg_facts_added']}")
        
        if self.stats['errors']:
            print(f"\nErrors ({len(self.stats['errors'])}):")
            for err in self.stats['errors']:
                print(f"  - {err}")
        
        print("=" * 60)


def main():
    # Find data directory
    data_dirs = [
        Path("D:/political-agent-sukumar/bkp/political-strategy-maker-v1/political-data"),
        Path("D:/political-agent-sukumar/political-strategy-maker/political-strategy-maker/political-data"),
    ]
    
    data_dir = None
    for d in data_dirs:
        if d.exists():
            data_dir = d
            break
    
    if not data_dir:
        print("ERROR: Data directory not found!")
        sys.exit(1)
    
    print(f"Data directory: {data_dir}")
    
    # Run ingestion
    ingester = BulkIngester(str(data_dir))
    ingester.initialize()
    ingester.process_all_files()
    ingester.print_summary()
    
    print("\n✅ Bulk ingestion complete!")
    print("Now restart the backend and test your queries.")


if __name__ == "__main__":
    main()

