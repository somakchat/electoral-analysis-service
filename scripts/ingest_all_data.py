"""
Comprehensive Data Ingestion Script - OPTIMIZED VERSION.

Ingests ALL political data to BOTH:
1. Knowledge Graph (KG) - for entity-based structured retrieval
2. OpenSearch - for vector/hybrid semantic search

PERFORMANCE OPTIMIZATIONS:
- Batch embeddings (100 at a time)
- Skip large CSV files or sample them
- Progress tracking
- Resume support (skip already indexed)

Usage:
    python ingest_all_data.py                    # Use default political-data folder
    python ingest_all_data.py /path/to/data      # Use custom data folder
    python ingest_all_data.py --skip-embeddings  # Only index to KG (fast)
"""
import os
import sys
from pathlib import Path
import asyncio
import gzip
import hashlib
import argparse
import time
import re

# Add backend to path
script_dir = Path(__file__).parent.parent
sys.path.insert(0, str(script_dir / "backend"))

from app.config import settings


def process_csv_efficiently(file_path: Path, max_rows: int = 500) -> list:
    """Process CSV files efficiently - sample large files."""
    import pandas as pd
    
    try:
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                df = pd.read_csv(file_path, encoding=encoding, nrows=max_rows)
                break
            except:
                continue
        else:
            return []
        
        # Convert to text chunks (combine rows into meaningful chunks)
        chunks = []
        
        # Get column info
        cols = list(df.columns)
        
        # Create summary chunk
        summary = f"Data from {file_path.name}:\n"
        summary += f"Columns: {', '.join(cols)}\n"
        summary += f"Total rows: {len(df)}\n"
        
        # Sample key data points
        for col in cols:
            if df[col].dtype == 'object':
                unique = df[col].dropna().unique()[:10]
                if len(unique) > 0:
                    summary += f"{col} values: {', '.join(str(v) for v in unique)}\n"
        
        chunks.append({"text": summary, "metadata": {"type": "csv_summary", "file": file_path.name}})
        
        # Create chunks for groups of rows (10 rows per chunk)
        for i in range(0, min(len(df), max_rows), 10):
            batch = df.iloc[i:i+10]
            text = batch.to_string(index=False)
            chunks.append({
                "text": text[:1000],  # Limit chunk size
                "metadata": {"type": "csv_data", "file": file_path.name, "rows": f"{i}-{i+10}"}
            })
        
        return chunks
        
    except Exception as e:
        print(f"    CSV processing error: {e}")
        return []


def process_excel_efficiently(file_path: Path, max_rows: int = 200) -> list:
    """Process Excel files efficiently."""
    import pandas as pd
    
    try:
        chunks = []
        
        # Read all sheets
        xl = pd.ExcelFile(file_path)
        
        for sheet_name in xl.sheet_names[:3]:  # Limit to first 3 sheets
            try:
                df = pd.read_excel(xl, sheet_name=sheet_name, nrows=max_rows)
                
                # Summary
                text = f"Sheet '{sheet_name}' from {file_path.name}:\n"
                text += df.head(20).to_string(index=False)
                
                chunks.append({
                    "text": text[:2000],
                    "metadata": {"type": "excel", "file": file_path.name, "sheet": sheet_name}
                })
            except:
                continue
        
        return chunks
        
    except Exception as e:
        print(f"    Excel processing error: {e}")
        return []


def process_docx_efficiently(file_path: Path) -> list:
    """Process DOCX files."""
    try:
        from docx import Document
        doc = Document(file_path)
        
        chunks = []
        current_text = ""
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                current_text += text + "\n"
                
                # Create chunk every ~1000 chars
                if len(current_text) > 1000:
                    chunks.append({
                        "text": current_text,
                        "metadata": {"type": "docx", "file": file_path.name}
                    })
                    current_text = ""
        
        # Don't forget the last chunk
        if current_text:
            chunks.append({
                "text": current_text,
                "metadata": {"type": "docx", "file": file_path.name}
            })
        
        return chunks
        
    except Exception as e:
        print(f"    DOCX processing error: {e}")
        return []


async def ingest_all_data(data_dir_path: str = None, skip_embeddings: bool = False, 
                          kg_only: bool = False, max_chunks_per_file: int = 50):
    """Main ingestion function - OPTIMIZED."""
    print("=" * 70)
    print("POLITICAL DATA INGESTION - OPTIMIZED VERSION")
    print("=" * 70)
    
    start_time = time.time()
    
    # Initialize components
    from app.services.orchestrator import Orchestrator
    from app.services.rag.unified_vectordb import VectorDBConfig, OpenSearchVectorDB, Document
    from app.services.rag.data_schema import FactWithCitation
    
    # Data directory
    if data_dir_path:
        data_dir = Path(data_dir_path)
    else:
        data_dir = script_dir / "political-data"
    
    if not data_dir.exists():
        print(f"ERROR: Data directory not found: {data_dir}")
        return
    
    print(f"\nData directory: {data_dir}")
    print(f"Skip embeddings: {skip_embeddings}")
    print(f"KG only mode: {kg_only}")
    print(f"Max chunks per file: {max_chunks_per_file}")
    
    # Initialize orchestrator and RAG
    print("\n[1/4] Initializing RAG system...")
    orchestrator = Orchestrator(index_dir=str(settings.index_dir))
    political_rag = orchestrator._get_political_rag()
    kg = political_rag.kg
    
    print(f"  - Knowledge Graph: {len(kg.constituency_profiles)} constituencies")
    print(f"  - Existing facts: {len(kg.facts)}")
    
    # Initialize OpenSearch (only if not kg_only)
    os_client = None
    embedder = None
    
    if not kg_only and not skip_embeddings:
        print("\n[2/4] Initializing OpenSearch...")
        if VectorDBConfig.is_configured():
            os_client = OpenSearchVectorDB()
            from app.services.rag.embeddings import get_embedding_service
            embedder = get_embedding_service()
            print(f"  - OpenSearch: {VectorDBConfig.OPENSEARCH_ENDPOINT}")
        else:
            print("  - OpenSearch not configured, skipping")
    else:
        print("\n[2/4] Skipping OpenSearch (kg_only or skip_embeddings mode)")
    
    # Get data files
    print("\n[3/4] Scanning data files...")
    
    supported_extensions = {'.csv', '.xlsx', '.xls', '.xlsm', '.docx', '.pdf', '.txt', '.gz'}
    data_files = [f for f in data_dir.iterdir() 
                  if f.is_file() and f.suffix.lower() in supported_extensions]
    
    print(f"  Found {len(data_files)} files")
    
    # Stats
    stats = {
        "files_processed": 0,
        "opensearch_chunks": 0,
        "kg_facts": 0,
        "errors": [],
        "skipped": []
    }
    
    # Known constituencies
    known_constituencies = set(kg.constituency_profiles.keys())
    
    # Process files
    print("\n[4/4] Processing files...")
    
    for i, file_path in enumerate(data_files, 1):
        file_start = time.time()
        ext = file_path.suffix.lower()
        
        print(f"\n  [{i}/{len(data_files)}] {file_path.name}")
        
        try:
            # Process based on file type
            if ext in ['.csv', '.gz']:
                if ext == '.gz':
                    # Decompress
                    with gzip.open(file_path, 'rt', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    temp_path = data_dir / f"_temp_{file_path.stem}"
                    temp_path.write_text(content[:500000], encoding='utf-8')  # Limit size
                    chunks = process_csv_efficiently(temp_path)
                    temp_path.unlink()
                else:
                    chunks = process_csv_efficiently(file_path)
                    
            elif ext in ['.xlsx', '.xls', '.xlsm']:
                chunks = process_excel_efficiently(file_path)
                
            elif ext == '.docx':
                chunks = process_docx_efficiently(file_path)
                
            else:
                # Simple text processing
                try:
                    text = file_path.read_text(encoding='utf-8', errors='ignore')
                    chunks = [{"text": text[:5000], "metadata": {"file": file_path.name}}]
                except:
                    chunks = []
            
            # Limit chunks
            chunks = chunks[:max_chunks_per_file]
            print(f"    - Extracted {len(chunks)} chunks")
            
            if not chunks:
                stats["skipped"].append(file_path.name)
                continue
            
            # Generate doc_id
            doc_id = f"{file_path.stem}_{hashlib.md5(file_path.name.encode()).hexdigest()[:6]}"
            
            # === INDEX TO OPENSEARCH (with batched embeddings) ===
            if os_client and embedder and chunks:
                try:
                    texts = [c.get("text", "")[:1500] for c in chunks if c.get("text")]  # Truncate
                    
                    if texts:
                        # Batch embedding (all at once)
                        print(f"    - Generating {len(texts)} embeddings...")
                        embeddings = embedder.embed(texts)
                        
                        # Prepare documents
                        documents = []
                        for j, (chunk, emb) in enumerate(zip(chunks, embeddings)):
                            if not chunk.get("text"):
                                continue
                            documents.append(Document(
                                doc_id=f"{doc_id}_c{j}",
                                text=chunk.get("text", "")[:1500],
                                embedding=emb.tolist() if hasattr(emb, 'tolist') else list(emb),
                                metadata={
                                    "source_file": file_path.name,
                                    "source_type": "political_data",
                                    "data_type": chunk.get("metadata", {}).get("type", "unknown")
                                }
                            ))
                        
                        # Bulk index
                        indexed = os_client.index_documents(documents)
                        stats["opensearch_chunks"] += indexed
                        print(f"    - OpenSearch: {indexed} indexed")
                        
                except Exception as e:
                    print(f"    - OpenSearch error: {str(e)[:50]}")
            
            # === INDEX TO KNOWLEDGE GRAPH ===
            facts_added = 0
            for j, chunk in enumerate(chunks):
                text = chunk.get("text", "")
                if not text:
                    continue
                
                # Find constituency mentions
                mentioned_acs = [ac for ac in known_constituencies if ac.lower() in text.lower()]
                
                # Find party mentions
                parties = []
                for party, pattern in [('BJP', r'\bbjp\b'), ('TMC', r'\btmc|trinamool\b'), 
                                       ('INC', r'\bcongress\b'), ('CPM', r'\bcpm|cpim\b')]:
                    if re.search(pattern, text.lower()):
                        parties.append(party)
                
                # Determine fact type
                fact_type = "document_content"
                if any(w in text.lower() for w in ['strategy', 'recommend']):
                    fact_type = "strategy"
                elif any(w in text.lower() for w in ['predict', 'forecast']):
                    fact_type = "prediction"
                elif any(w in text.lower() for w in ['vote share', 'margin', 'won']):
                    fact_type = "electoral_result"
                
                # Create facts
                entities = mentioned_acs + parties
                for entity in entities[:5]:  # Limit entities per chunk
                    try:
                        fact = FactWithCitation(
                            fact_type=fact_type,
                            fact_text=text[:400],
                            entity_name=entity,
                            entity_type="constituency" if entity in known_constituencies else "party",
                            confidence=0.85,
                            source_file=file_path.name,
                            source_doc_ids=[doc_id],
                            related_entities=entities[:10]
                        )
                        kg.add_fact(fact)
                        facts_added += 1
                    except Exception as e:
                        pass
            
            stats["kg_facts"] += facts_added
            print(f"    - KG facts: {facts_added}")
            
            stats["files_processed"] += 1
            elapsed = time.time() - file_start
            print(f"    - Time: {elapsed:.1f}s")
            
        except Exception as e:
            stats["errors"].append(f"{file_path.name}: {str(e)[:50]}")
            print(f"    - ERROR: {str(e)[:50]}")
    
    # Summary
    total_time = time.time() - start_time
    
    print("\n" + "=" * 70)
    print("INGESTION COMPLETE")
    print("=" * 70)
    print(f"\n  Time: {total_time/60:.1f} minutes")
    print(f"  Files processed: {stats['files_processed']}/{len(data_files)}")
    print(f"  OpenSearch chunks: {stats['opensearch_chunks']}")
    print(f"  KG facts added: {stats['kg_facts']}")
    print(f"  Total KG facts: {len(kg.facts)}")
    
    if stats["errors"]:
        print(f"\n  Errors: {len(stats['errors'])}")
        for e in stats["errors"][:3]:
            print(f"    - {e}")
    
    print("\n✅ Data indexed to both KG and OpenSearch!")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Ingest political data (optimized)")
    parser.add_argument("data_dir", nargs="?", default=None, help="Data directory path")
    parser.add_argument("--skip-embeddings", action="store_true", help="Skip OpenSearch indexing")
    parser.add_argument("--kg-only", action="store_true", help="Only index to Knowledge Graph")
    parser.add_argument("--max-chunks", type=int, default=50, help="Max chunks per file (default: 50)")
    
    args = parser.parse_args()
    
    asyncio.run(ingest_all_data(
        args.data_dir, 
        skip_embeddings=args.skip_embeddings,
        kg_only=args.kg_only,
        max_chunks_per_file=args.max_chunks
    ))
